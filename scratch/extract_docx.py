#!/usr/bin/env python3
import sys
from pathlib import Path
from docx import Document

def main():
    if len(sys.argv) < 2:
        print("Usage: extract_docx.py <path_to_docx>")
        sys.exit(1)
        
    docx_path = Path(sys.argv[1])
    if not docx_path.exists():
        print(f"Error: file {docx_path} does not exist")
        sys.exit(1)
        
    doc = Document(str(docx_path))
    print(f"--- Document: {docx_path.name} ---")
    
    # Extract paragraphs
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    print(f"Found {len(paragraphs)} paragraphs.")
    
    # Extract tables
    tables_text = []
    for i, table in enumerate(doc.tables):
        rows_data = []
        for row in table.rows:
            cells_data = [cell.text.strip() for cell in row.cells]
            # remove duplicates in adjacent cells due to merged cells
            cleaned_cells = []
            for cell in cells_data:
                if not cleaned_cells or cell != cleaned_cells[-1]:
                    cleaned_cells.append(cell)
            rows_data.append(" | ".join(cleaned_cells))
        tables_text.append(f"--- Table {i+1} ---\n" + "\n".join(rows_data))
    
    print(f"Found {len(doc.tables)} tables.")
    
    output_text = "\n\n".join(paragraphs) + "\n\n" + "\n\n".join(tables_text)
    
    # Print first 2000 chars and save full text to a temp file next to the script
    print("\n--- SAMPLE CONTENT (First 2000 chars) ---")
    print(output_text[:2000])
    print("------------------------------------------")
    
    txt_path = docx_path.with_suffix(".txt")
    txt_path.write_text(output_text, encoding="utf-8")
    print(f"Full text extracted and saved to: {txt_path}")

if __name__ == "__main__":
    main()
