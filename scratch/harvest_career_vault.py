#!/usr/bin/env python3
"""
harvest_career_vault.py — Career Brain ETL Harvesting & Processing Engine
========================================================================
Reads the structured consolidated_vault_audit_raw.json registry, extracts
text natively from all high-value formats (docx, pdf, xlsx, json, html, csv),
and writes standard plain-text versions into active raw_docs/ subdirectories
preserving strict lineage.
"""

import os
import re
import zipfile
import json
import xml.etree.ElementTree as ET
from pathlib import Path
from datetime import datetime

# Try dependencies
try:
    from docx import Document
except ImportError:
    Document = None

try:
    from pypdf import PdfReader
except ImportError:
    PdfReader = None

# Paths
BASE = Path("/Users/okgoogle13/Projects/Career Brain")
RAW_DOCS = BASE / "raw_docs"
AUDIT_JSON_PATH = BASE / "scratch" / "consolidated_vault_audit_raw.json"
HARVEST_LOG = BASE / "output" / "harvest_run.log"

# Category mappings to raw_docs
ROUTING = {
    "Nishant's authentic personal career history/application document.": "resumes",
    "Master consolidated CV containing high-fidelity tailored items.": "resumes",
    "Nishant's direct selection criteria STAR narratives.": "ksc",
    "AI prompting templates, Custom Gem instructions, or agent manual.": "knowledge",
    "Australian community sector-specific keywords, skills translations, or inclusive language guidelines.": "knowledge",
    "Standard guidelines on LGBTIQA+ inclusive language in Australia.": "knowledge",
    "Rosetta Stone translation taxonomy mapping guides.": "knowledge",
    "Reference resume example for sector benchmarking.": "resumes",
    "Reference document": "knowledge"
}

def sanitize_name(name: str) -> str:
    # Remove extension
    stem = Path(name).stem
    # Replace spaces and special characters with underscores
    stem = re.sub(r"[^\w\s\-\.]", "_", stem)
    stem = re.sub(r"\s+", "_", stem)
    return stem

def extract_xlsx(path: Path) -> str:
    try:
        with zipfile.ZipFile(str(path)) as z:
            if "xl/sharedStrings.xml" in z.namelist():
                xml_content = z.read("xl/sharedStrings.xml")
                root = ET.fromstring(xml_content)
                texts = [elem.text for elem in root.iter() if elem.tag.endswith('t') and elem.text]
                return "\n".join(texts)
            else:
                return "[XLSX has no shared strings]"
    except Exception as e:
        return f"[Error reading XLSX: {e}]"

def extract_html(path: Path) -> str:
    try:
        content = path.read_text(encoding="utf-8", errors="replace")
        content = re.sub(r'<(script|style).*?>.*?</\1>', '', content, flags=re.DOTALL | re.IGNORECASE)
        clean_text = re.sub(r'<.*?>', ' ', content)
        clean_text = re.sub(r'\s+', ' ', clean_text).strip()
        return clean_text
    except Exception as e:
        return f"[Error reading HTML: {e}]"

def extract_docx(path: Path) -> str:
    if Document is None:
        raise ImportError("python-docx not installed")
    doc = Document(str(path))
    paragraphs = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            row_text = []
            for cell in row.cells:
                row_text.append(cell.text.strip())
            paragraphs.append(" | ".join(row_text))
    return "\n".join(paragraphs)

def extract_pdf(path: Path) -> str:
    if PdfReader is None:
        raise ImportError("pypdf not installed")
    reader = PdfReader(str(path))
    pages = []
    for page in reader.pages:
        pages.append(page.extract_text() or "")
    return "\n\n".join(pages)

def extract_content(path: Path) -> str:
    ext = path.suffix.lower()
    if ext in [".txt", ".md", ".csv"]:
        return path.read_text(encoding="utf-8", errors="replace")
    elif ext == ".docx":
        return extract_docx(path)
    elif ext == ".pdf":
        return extract_pdf(path)
    elif ext == ".xlsx":
        return extract_xlsx(path)
    elif ext == ".html":
        return extract_html(path)
    elif ext in [".json", ".jsonl"]:
        # Pretty print JSON cleanly
        try:
            data = json.loads(path.read_text(encoding="utf-8", errors="replace"))
            return json.dumps(data, indent=2)
        except Exception:
            return path.read_text(encoding="utf-8", errors="replace")
    else:
        return ""

def main():
    print("=" * 60)
    print("CAREER BRAIN - DATA HARVESTING & PROCESSING WORKFLOW")
    print("=" * 60)
    
    if not AUDIT_JSON_PATH.exists():
        print(f"Error: Audit database JSON not found at {AUDIT_JSON_PATH}.")
        print("Please run scratch/consolidated_vault_audit.py first.")
        return

    # Load registry
    registry = json.loads(AUDIT_JSON_PATH.read_text(encoding="utf-8"))
    
    # Filter to high-value unique files
    high_value_items = [f for f in registry if f["value"] in ["Very High", "High"]]
    
    print(f"Found {len(high_value_items)} high-value unique files eligible for harvesting.")
    
    stats = {
        "resumes": 0,
        "cover_letters": 0,
        "ksc": 0,
        "knowledge": 0,
        "references": 0,
        "skipped": 0,
        "errors": 0
    }
    
    log_entries = []
    
    # Process each
    for item in high_value_items:
        name = item["name"]
        folder_slug = item["folder"]
        purpose = item["purpose"]
        
        # Determine original full path
        original_path = BASE / folder_slug / name
        if not original_path.exists():
            # Try absolute path just in case
            original_path = Path(folder_slug) / name
            if not original_path.exists():
                print(f"Warning: File {name} not found at {original_path}. Skipping.")
                stats["skipped"] += 1
                continue
                
        # Resolve raw_docs routing
        route_category = ROUTING.get(purpose, None)
        
        # Backup custom routing logic
        if not route_category:
            lower_name = name.lower()
            if "resume" in lower_name or "cv" in lower_name:
                route_category = "resumes"
            elif "cover_letter" in lower_name or "cover letter" in lower_name:
                route_category = "cover_letters"
            elif "ksc" in lower_name or "selection criteria" in lower_name or "statement of claims" in lower_name:
                route_category = "ksc"
            elif "reference" in lower_name:
                route_category = "references"
            else:
                route_category = "knowledge"
                
        # We do not want to harvest files that are already inside raw_docs/ folder
        if folder_slug.startswith("raw_docs"):
            print(f"Skipping active source file: {name} (Already in {folder_slug})")
            stats["skipped"] += 1
            continue
            
        target_dir = RAW_DOCS / route_category
        target_dir.mkdir(parents=True, exist_ok=True)
        
        # Build strict lineage filename to prevent collisions and preserve origin
        folder_clean = sanitize_name(folder_slug)
        file_clean = sanitize_name(name)
        
        harvested_name = f"harvested_{folder_clean}__{file_clean}.txt"
        target_path = target_dir / harvested_name
        
        print(f"Harvesting: {folder_slug}/{name} → raw_docs/{route_category}/{harvested_name}")
        
        try:
            # Extract content natively
            text_content = extract_content(original_path)
            
            if not text_content or len(text_content.strip()) < 20:
                print(f"  [ERROR] Empty or too short text extracted from {name}")
                stats["errors"] += 1
                log_entries.append(f"ERROR | {name} | Extracted text empty or too short.")
                continue
                
            # Construct strict header block for pipeline tracing
            header = (
                f"=== HARVESTED SOURCE LINEAGE ===\n"
                f"Original Directory : {folder_slug}\n"
                f"Original Filename  : {name}\n"
                f"Original Purpose   : {purpose}\n"
                f"Harvested Runtime  : {datetime.now().isoformat()}\n"
                f"{'='*60}\n\n"
            )
            
            target_path.write_text(header + text_content.strip(), encoding="utf-8")
            stats[route_category] += 1
            log_entries.append(f"SUCCESS | {folder_slug}/{name} → raw_docs/{route_category}/{harvested_name}")
            
        except Exception as e:
            print(f"  [ERROR] Failed to extract {name}: {e}")
            stats["errors"] += 1
            log_entries.append(f"ERROR | {name} | Extraction failed: {e}")
            
    # Write summary log
    summary = [
        "=== CAREER BRAIN HARVESTING WORKFLOW RUN LOG ===",
        f"Timestamp: {datetime.now().isoformat()}",
        f"Total Ingested Files: {sum(stats.values()) - stats['skipped'] - stats['errors']}",
        f"  Resumes       : {stats['resumes']}",
        f"  Cover Letters : {stats['cover_letters']}",
        f"  KSC/STAR      : {stats['ksc']}",
        f"  References    : {stats['references']}",
        f"  Knowledge     : {stats['knowledge']}",
        f"Errors          : {stats['errors']}",
        f"Skipped (Active): {stats['skipped']}",
        "================================================",
        ""
    ] + log_entries
    
    HARVEST_LOG.parent.mkdir(parents=True, exist_ok=True)
    HARVEST_LOG.write_text("\n".join(summary) + "\n", encoding="utf-8")
    
    print("\n" + "=" * 60)
    print("HARVESTING RUN COMPLETE")
    print("=" * 60)
    print(f"  ✓ Resumes Harvested       : {stats['resumes']}")
    print(f"  ✓ Cover Letters Harvested : {stats['cover_letters']}")
    print(f"  ✓ KSC Narratives Harvested: {stats['ksc']}")
    print(f"  ✓ References Harvested    : {stats['references']}")
    print(f"  ✓ Knowledge/Glossaries    : {stats['knowledge']}")
    print(f"  ✗ Errors Encountered     : {stats['errors']}")
    print(f"  - Skipped Active Sources  : {stats['skipped']}")
    print(f"Detailed run log saved to   : {HARVEST_LOG}")
    print("=" * 60)

if __name__ == "__main__":
    main()
