#!/usr/bin/env python3
import os
import re
import zipfile
import hashlib
import json
import xml.etree.ElementTree as ET
from pathlib import Path

# Try dependencies
try:
    from docx import Document
except ImportError:
    Document = None

try:
    from pypdf import PdfReader
except ImportError:
    PdfReader = None

# Comprehensive set of directories to audit
DIRS = {
    "raw_docs/knowledge": Path("/Users/okgoogle13/Projects/Career Brain/raw_docs/knowledge"),
    "CLEANUP/AI Knowledge Files": Path("/Users/okgoogle13/Projects/Career Brain/CLEANUP/AI Knowledge Files"),
    "CLEANUP/Bestpractice": Path("/Users/okgoogle13/Projects/Career Brain/CLEANUP/Bestpractice"),
    "CLEANUP/CUSTOM GEM REFERENCE STUFF": Path("/Users/okgoogle13/Projects/Career Brain/CLEANUP/CUSTOM GEM REFERENCE STUFF"),
    "CLEANUP/Chromebook Downloads": Path("/Users/okgoogle13/Projects/Career Brain/CLEANUP/Chromebook Downloads"),
    "CLEANUP/Key Selection Criteria Responses": Path("/Users/okgoogle13/Projects/Career Brain/CLEANUP/Key Selection Criteria Responses"),
    "CLEANUP/Old cover letters": Path("/Users/okgoogle13/Projects/Career Brain/CLEANUP/Old cover letters"),
    "CLEANUP/Old resumes ": Path("/Users/okgoogle13/Projects/Career Brain/CLEANUP/Old resumes "),
    "CLEANUP/References": Path("/Users/okgoogle13/Projects/Career Brain/CLEANUP/References"),
    "CLEANUP/Resume templates ": Path("/Users/okgoogle13/Projects/Career Brain/CLEANUP/Resume templates ")
}

REPORT_PATH = Path("/Users/okgoogle13/.gemini/antigravity-ide/brain/0666d05e-cac5-45ef-b06f-a8ff6cb41b76/consolidated_vault_audit_results.md")

def get_hash(path: Path) -> str:
    try:
        h = hashlib.md5()
        h.update(path.read_bytes())
        return h.hexdigest()
    except Exception as e:
        return f"error_{e}"

def extract_xlsx(path: Path) -> str:
    try:
        with zipfile.ZipFile(str(path)) as z:
            if "xl/sharedStrings.xml" in z.namelist():
                xml_content = z.read("xl/sharedStrings.xml")
                root = ET.fromstring(xml_content)
                texts = [elem.text for elem in root.iter() if elem.tag.endswith('t') and elem.text]
                return "\n".join(texts)[:1500]
            else:
                return "[XLSX has no shared strings]"
    except Exception as e:
        return f"[Error reading XLSX: {e}]"

def extract_html(path: Path) -> str:
    try:
        content = path.read_text(encoding="utf-8", errors="replace")
        content = re.sub(r'<(script|style).*?>.*?</\1>', '', content, flags=re.DOTALL | re.IGNORECASE)
        clean_text = re.sub(r'<.*?>', ' ', content)
        clean_text = re.sub(r'\s+', ' ', clean_text).strip()
        return clean_text[:1500]
    except Exception as e:
        return f"[Error reading HTML: {e}]"

def extract_json(path: Path) -> str:
    try:
        content = path.read_text(encoding="utf-8", errors="replace")
        data = json.loads(content)
        return json.dumps(data, indent=2)[:1500]
    except Exception as e:
        return f"[Error reading JSON: {e}]"

def extract_snippet(path: Path) -> str:
    ext = path.suffix.lower()
    if ext in [".txt", ".md", ".csv"]:
        try:
            content = path.read_text(encoding="utf-8", errors="replace").strip()
            return content[:1500]
        except Exception as e:
            return f"[Error reading text: {e}]"
    elif ext == ".docx":
        if Document is None:
            return "[docx library missing]"
        try:
            doc = Document(str(path))
            text = "\n".join([p.text for p in doc.paragraphs[:15]])
            for t in doc.tables[:3]:
                for row in t.rows[:3]:
                    for cell in row.cells[:3]:
                        text += f" | {cell.text}"
            return text[:1500]
        except Exception as e:
            return f"[Error reading docx: {e}]"
    elif ext == ".pdf":
        if PdfReader is None:
            return "[pdf library missing]"
        try:
            reader = PdfReader(str(path))
            if not reader.pages:
                return "[Empty PDF]"
            first_page_text = reader.pages[0].extract_text() or ""
            return first_page_text[:1500]
        except Exception as e:
            return f"[Error reading pdf: {e}]"
    elif ext == ".xlsx":
        return extract_xlsx(path)
    elif ext == ".html":
        return extract_html(path)
    elif ext in [".json", ".jsonl"]:
        return extract_json(path)
    elif ext in [".png", ".jpg", ".jpeg", ".gif"]:
        return f"[Image file: {path.name}]"
    elif ext in [".zip", ".tar", ".gz"]:
        return f"[Archive file: {path.name}]"
    else:
        return f"[File preview not available for extension: {ext}]"

def evaluate_file(f, all_files):
    name = f["name"]
    ext = f["ext"]
    size_kb = round(f["size"] / 1024, 2)
    snippet = f["snippet"]
    is_dup = f["is_duplicate"]
    dup_of = f["duplicate_of"]
    folder = f["folder"]
    
    # 1. Byte duplicates are Very Low
    if is_dup:
        return {
            "name": name,
            "folder": folder,
            "ext": ext,
            "size_kb": size_kb,
            "purpose": f"Identical clone of {dup_of}",
            "value": "Very Low",
            "action": "Delete Duplicate",
            "snippet": f"[Duplicate of {dup_of}]"
        }
        
    lower_name = name.lower()
    lower_snippet = snippet.lower()
    
    # Temp files from MS Word
    if name.startswith("~$"):
        return {
            "name": name,
            "folder": folder,
            "ext": ext,
            "size_kb": size_kb,
            "purpose": "MS Word temporary file remnant.",
            "value": "Very Low",
            "action": "Delete Temp File",
            "snippet": "[Temp File]"
        }
        
    # Defaults
    purpose = "Reference document"
    value = "Med"
    action = "Keep"

    # Nishant's direct personal application/resume files
    if "nishant" in lower_name or "resume" in lower_name or "cover_letter" in lower_name or "ksc_responses" in lower_name or "cv" in lower_name or "cover letter" in lower_name:
        if "nishant_dougall" in lower_name or "nishant dougall" in lower_name or "nish_dougall" in lower_name or "peer resume" in lower_snippet or "peer role" in lower_name or "peer worker" in lower_name:
            purpose = "Nishant's authentic personal career history/application document."
            value = "Very High"
            action = "Harvest Content & Archive Source"
        elif "example" in lower_name or "sample" in lower_name or "template" in lower_name or "boring" in lower_name or "khaki" in lower_name or "stripe" in lower_name:
            purpose = "Reference template or styling layout format."
            value = "Med"
            action = "Keep in archive / reference"
            
    if "final combined resumes" in lower_name or "master resume with subtypes" in lower_name or "master_resume" in lower_name:
        purpose = "Master consolidated CV containing high-fidelity tailored items."
        value = "Very High"
        action = "Keep & sync with master resume"

    # Selection Criteria Responses
    if "selection criteria" in lower_name or "ksc" in lower_name or "statement of claims" in lower_name or "key criteria" in lower_name:
        if "nishant" in lower_name or "dougall" in lower_name or "iap ksc" in lower_name or "flat out ksc" in lower_name or "thh peer worker" in lower_name:
            purpose = "Nishant's direct selection criteria STAR narratives."
            value = "Very High"
            action = "Harvest Narratives & Archive Source"

    # Prompting Guides and AI Agent Systems
    if "prompt" in lower_name or "customgem" in lower_name or "gemini" in lower_name or "manual" in lower_name or "solution design" in lower_name:
        purpose = "AI prompting templates, Custom Gem instructions, or agent manual."
        value = "Very High"
        action = "Keep & integrate into system prompts"

    # Keywords, Glossaries, and Taxonomies
    if "keyword" in lower_name or "taxonomy" in lower_name or "glossary" in lower_name or "definitions" in lower_name or "inclusive-language" in lower_name or "lgbtiqa" in lower_name or "terminology" in lower_name:
        purpose = "Australian community sector-specific keywords, skills translations, or inclusive language guidelines."
        value = "High"
        action = "Extract terms & sync with skills_and_taxonomy.json"
        
    # Non-career personal files in Chromebook Downloads (e.g. travel itineraries, bills, invoices, receipts, medical super releases)
    personal_docs = ["itinerary", "travel planner", "american express", "ovo energy", "superannuation", "bmw", "billing account", "invoice", "receipt", "statutory-declaration", "change-of-sex", "change of name", "declarations", "declara", "whatsapp chat", "pxl_", "screenshot", "material-theme"]
    if any(p in lower_name for p in personal_docs):
        purpose = "Personal non-career document (medical, finance, travel, administrative)."
        value = "Low"
        action = "Move to secure external backup / Archive"

    # Third-party generic advice
    if "seek" in lower_name or "indeed" in lower_name or "hays" in lower_name or "atlassian" in lower_name or "ethicaljobs" in lower_name or "jobscan" in lower_name or "margins" in lower_name or "cliche" in lower_name:
        if value != "Very High" and value != "High":
            purpose = "Standard third-party generic job application guidance article."
            value = "Low"
            action = "Archive to third party advice folder"

    # Format duplicates: if text copy of a rich file
    name_no_ext = Path(name).stem
    if ext.lower() in [".txt", ".csv"]:
        for other in all_files:
            if other["name"] != name and Path(other["name"]).stem == name_no_ext:
                if other["ext"].lower() in [".md", ".docx", ".pdf"]:
                    purpose = f"Flat text representation of {other['ext']} file."
                    value = "Very Low"
                    action = "Delete Redundant Format"
                    break

    return {
        "name": name,
        "folder": folder,
        "ext": ext,
        "size_kb": size_kb,
        "purpose": purpose,
        "value": value,
        "action": action,
        "snippet": snippet[:150].replace('\n', ' ') + "..."
    }

def main():
    print("=" * 60)
    print("CAREER BRAIN - CONSOLIDATED VAULT SCANNER")
    print("=" * 60)
    
    all_files_info = []
    seen_hashes = {}
    
    for folder_name, folder_path in DIRS.items():
        if not folder_path.exists():
            print(f"Directory {folder_path} does not exist. Skipping.")
            continue
            
        print(f"Scanning folder: {folder_name}...")
        for root, subdirs, files in os.walk(folder_path):
            for filename in sorted(files):
                if filename.startswith("."):
                    continue
                filepath = Path(root) / filename
                
                fhash = get_hash(filepath)
                size = filepath.stat().st_size
                snippet = extract_snippet(filepath)
                
                is_duplicate = False
                duplicate_of = None
                
                if fhash.startswith("error"):
                    pass
                elif fhash in seen_hashes:
                    is_duplicate = True
                    duplicate_of = seen_hashes[fhash]
                else:
                    seen_hashes[fhash] = f"{folder_name}/{filename}"
                    
                all_files_info.append({
                    "name": filename,
                    "folder": folder_name,
                    "ext": filepath.suffix.lower(),
                    "size": size,
                    "hash": fhash,
                    "is_duplicate": is_duplicate,
                    "duplicate_of": duplicate_of,
                    "snippet": snippet
                })
                
    print(f"Scanned a total of {len(all_files_info)} raw files.")
    
    # Process files
    analyzed = []
    for f in all_files_info:
        analyzed.append(evaluate_file(f, all_files_info))
        
    # Sort
    value_order = {"Very High": 0, "High": 1, "Med": 2, "Low": 3, "Very Low": 4}
    analyzed.sort(key=lambda x: (value_order.get(x["value"], 5), x["folder"], x["name"]))
    
    # Generate stats
    stats = {"Very High": 0, "High": 0, "Med": 0, "Low": 0, "Very Low": 0}
    folder_stats = {}
    for item in analyzed:
        stats[item["value"]] += 1
        folder_stats.setdefault(item["folder"], {}).setdefault(item["value"], 0)
        folder_stats[item["folder"]][item["value"]] += 1
        
    print("\nValue ratings breakdown:")
    for val, count in stats.items():
        print(f"  {val}: {count}")
        
    # Save a JSON file with full processed audit details
    json_path = Path("/Users/okgoogle13/Projects/Career Brain/scratch/consolidated_vault_audit_raw.json")
    json_path.write_text(json.dumps(analyzed, indent=2), encoding="utf-8")
    print(f"Saved full structured audit registry JSON to: {json_path}")

    # Generate Markdown report
    md = []
    md.append("# **Career Brain Comprehensive Vault Consolidation & Audit Report**\n")
    md.append(f"Generated: consolidated_vault_audit.py at local system runtime.")
    md.append("\nThis report represents a complete and comprehensive audit of **all 10 directories** containing active reference material, backups, and legacy prototypes inside the Career Brain project:")
    for i, folder in enumerate(DIRS.keys(), 1):
        md.append(f"{i}. `{folder}`")
        
    md.append("\n## **Audit Summary & Value Distribution**\n")
    md.append("| Value Rating | Count | Description | Primary Pipeline Action |")
    md.append("| --- | --- | --- | --- |")
    md.append(f"| 🟢 **Very High** | {stats['Very High']} | Gold source personal resumes, custom KSC STAR narratives, prompt design sheets. | **Extract and integrate into the main pipeline.** |")
    md.append(f"| 🔵 **High** | {stats['High']} | Community services vocabulary lists, inclusive language guides, skill taxonomies. | **Ingest and sync with `skills_and_taxonomy.json`.** |")
    md.append(f"| 🟡 **Med** | {stats['Med']} | Standard templates, project blueprints, checklists, config schemas. | **Keep in historical backup directory.** |")
    md.append(f"| 🟠 **Low** | {stats['Low']} | Generic jobsite articles or administrative files (itineraries, invoices, bills). | **Move out of Career Brain project workspace.** |")
    md.append(f"| 🔴 **Very Low** | {stats['Very Low']} | Byte-for-byte exact duplicate files, MS Word temp file remnants. | **Purge / Delete safely to reclaim space.** |")
    
    md.append("\n## **Directory Inventory & Value Distribution**\n")
    md.append("| Directory | Total Files | Very High | High | Med | Low | Very Low |")
    md.append("| --- | --- | --- | --- | --- | --- | --- |")
    for fld in sorted(DIRS.keys()):
        f_stats = folder_stats.get(fld, {})
        tot = sum(f_stats.values())
        md.append(f"| `{fld}` | {tot} | {f_stats.get('Very High',0)} | {f_stats.get('High',0)} | {f_stats.get('Med',0)} | {f_stats.get('Low',0)} | {f_stats.get('Very Low',0)} |")
        
    md.append("\n## **Full File Registry & Action Ledger**\n")
    md.append("Below is the complete semantic inventory of all analyzed files, ordered from highest value to lowest value.\n")
    md.append("| File Name | Folder | Size (KB) | Value | Purpose | Recommended Action | Snippet Preview |")
    md.append("| --- | --- | --- | --- | --- | --- | --- |")
    for item in analyzed:
        clean_snippet = item["snippet"].replace("|", "\\|")
        md.append(f"| `{item['name']}` | `{item['folder']}` | {item['size_kb']} KB | **{item['value']}** | {item['purpose']} | `{item['action']}` | *{clean_snippet}* |")
        
    REPORT_PATH.parent.mkdir(parents=True, exist_ok=True)
    REPORT_PATH.write_text("\n".join(md), encoding="utf-8")
    print(f"Audit report generated successfully at: {REPORT_PATH}")

if __name__ == "__main__":
    main()
