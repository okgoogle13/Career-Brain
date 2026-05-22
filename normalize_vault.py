#!/usr/bin/env python3
"""
normalize_vault.py — Career Brain Pipeline: Phase 1
====================================================
Converts all source documents in raw_docs/ into plain .txt files
in normalized_vault/. No summarizing, redacting, or semantic processing.
Pure string extraction only.

Supported formats: .docx, .doc (OLE2 via olefile), .pdf, .txt, .md, .csv
"""

import os
import re
import sys
import hashlib
import logging
from pathlib import Path
from datetime import datetime

# ─── Dependencies ───────────────────────────────────────────────────────────
try:
    from docx import Document as DocxDocument
except ImportError:
    DocxDocument = None

try:
    from pypdf import PdfReader
except ImportError:
    PdfReader = None

try:
    import olefile
    import chardet
    OLEFILE_AVAILABLE = True
except ImportError:
    OLEFILE_AVAILABLE = False

import subprocess
ANTIWORD = None
for candidate in ["/opt/homebrew/bin/antiword", "/usr/local/bin/antiword"]:
    if Path(candidate).exists():
        ANTIWORD = candidate
        break
if ANTIWORD is None:
    # try PATH
    try:
        result = subprocess.run(["which", "antiword"], capture_output=True, text=True)
        if result.returncode == 0:
            ANTIWORD = result.stdout.strip()
    except Exception:
        pass

# ─── Paths ───────────────────────────────────────────────────────────────────
BASE          = Path(__file__).parent
RAW_DOCS      = BASE / "raw_docs"
VAULT         = BASE / "normalized_vault"
OUTPUT        = BASE / "output"
ERROR_LOG     = OUTPUT / "parsing_errors.log"

VAULT.mkdir(parents=True, exist_ok=True)
OUTPUT.mkdir(parents=True, exist_ok=True)

# ─── Logging ─────────────────────────────────────────────────────────────────
logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s  %(levelname)-8s  %(message)s",
    handlers=[
        logging.StreamHandler(sys.stdout),
    ]
)
log = logging.getLogger("normalize_vault")

error_entries = []

def log_error(filename: str, reason: str):
    entry = f"[{datetime.now().isoformat()}] ERROR | {filename} | {reason}"
    error_entries.append(entry)
    log.warning(f"PARSE ERROR: {filename} → {reason}")

# ─── Extraction Functions ─────────────────────────────────────────────────────

def extract_docx(path: Path) -> str:
    """Extract text from .docx via python-docx."""
    if DocxDocument is None:
        raise ImportError("python-docx not installed")
    doc = DocxDocument(str(path))
    paragraphs = [p.text for p in doc.paragraphs]
    # Also extract from tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                paragraphs.append(cell.text)
    return "\n".join(paragraphs)


def extract_doc_ole(path: Path) -> str:
    """
    Fallback: extract text from legacy .doc (OLE2/Word97-2003) via olefile.
    Used only when antiword is unavailable.
    """
    if not OLEFILE_AVAILABLE:
        raise ImportError("olefile not installed")

    with olefile.OleFileIO(str(path)) as ole:
        if ole.exists("WordDocument"):
            data = ole.openstream("WordDocument").read()
        else:
            raise ValueError("No WordDocument stream found in OLE file")

    text_parts = []
    try:
        decoded = data.decode("utf-16-le", errors="ignore")
        text_parts = [c for c in decoded if c.isprintable() or c in "\n\r\t "]
    except Exception:
        pass

    if not text_parts or len("".join(text_parts).strip()) < 50:
        try:
            decoded = data.decode("latin-1", errors="ignore")
            text_parts = [c for c in decoded if c.isprintable() or c in "\n\r\t "]
        except Exception:
            pass

    raw_text = "".join(text_parts)
    raw_text = re.sub(r"\r\n|\r", "\n", raw_text)
    raw_text = re.sub(r"[ \t]{2,}", " ", raw_text)
    raw_text = re.sub(r"\n{4,}", "\n\n\n", raw_text)
    return raw_text.strip()


def extract_doc(path: Path) -> str:
    """
    Primary .doc extractor: uses antiword (brew install antiword) for clean,
    reliable Word97-2003 text extraction. Falls back to olefile byte parsing.
    """
    # 1. Try antiword (best quality)
    if ANTIWORD:
        try:
            result = subprocess.run(
                [ANTIWORD, "-w", "0", str(path)],
                capture_output=True,
                timeout=30
            )
            if result.returncode == 0 and result.stdout:
                text = result.stdout.decode("utf-8", errors="replace")
                if len(text.strip()) > 20:
                    return text.strip()
            # antiword returned empty or error — fall through
            stderr = result.stderr.decode("utf-8", errors="ignore")
            if stderr:
                log.debug(f"antiword stderr for {path.name}: {stderr.strip()}")
        except subprocess.TimeoutExpired:
            log.warning(f"antiword timed out on {path.name}")
        except Exception as e:
            log.debug(f"antiword failed on {path.name}: {e}")

    # 2. Fall back to olefile
    return extract_doc_ole(path)


def extract_pdf(path: Path) -> str:
    """Extract text from .pdf via pypdf."""
    if PdfReader is None:
        raise ImportError("pypdf not installed")
    reader = PdfReader(str(path))
    pages = []
    for page in reader.pages:
        pages.append(page.extract_text() or "")
    return "\n\n".join(pages)


def extract_text(path: Path) -> str:
    """Read .txt, .md, .csv files directly."""
    try:
        return path.read_text(encoding="utf-8")
    except UnicodeDecodeError:
        detected = chardet.detect(path.read_bytes()) if OLEFILE_AVAILABLE else {}
        enc = detected.get("encoding") or "latin-1"
        return path.read_text(encoding=enc, errors="replace")


# ─── Router ──────────────────────────────────────────────────────────────────

EXTRACTORS = {
    ".docx": extract_docx,
    ".doc":  extract_doc,       # antiword primary, olefile fallback
    ".pdf":  extract_pdf,
    ".txt":  extract_text,
    ".md":   extract_text,
    ".csv":  extract_text,
}

SKIP_EXTENSIONS = {".zip", ".xlsx", ".png", ".jpg", ".jpeg", ".mp4",
                   ".dotx", ".dotm", ".json", ".html", ".psl", ".textClipping",
                   ".DOCX"}  # uppercase handled by .lower() below


def file_md5(path: Path) -> str:
    h = hashlib.md5()
    h.update(path.read_bytes())
    return h.hexdigest()


def safe_stem(src_dir_name: str, filename: str) -> str:
    """Create a vault filename: {source_category}_{original_stem}.txt"""
    stem = Path(filename).stem
    # Sanitize: replace special chars except alphanumeric, underscore, hyphen, space, dot
    stem = re.sub(r"[^\w\s\-\.]", "_", stem)
    stem = re.sub(r"\s+", "_", stem)
    return f"{src_dir_name}__{stem}.txt"


# ─── Main Processing Loop ────────────────────────────────────────────────────

def process_directory(src_dir: Path, category: str, seen_hashes: dict) -> dict:
    """Process all files in src_dir, return stats."""
    stats = {"processed": 0, "skipped_ext": 0, "skipped_dup": 0, "errors": 0, "files": []}

    if not src_dir.exists():
        log.warning(f"Directory not found: {src_dir}")
        return stats

    files = sorted(src_dir.iterdir())
    log.info(f"\n{'─'*60}")
    log.info(f"Processing: {src_dir.name}  ({len(files)} files)")
    log.info(f"{'─'*60}")

    for f in files:
        if not f.is_file():
            continue

        ext = f.suffix.lower()

        # Skip unsupported extensions
        if ext in SKIP_EXTENSIONS or ext.upper() in SKIP_EXTENSIONS or ext not in EXTRACTORS:
            log.debug(f"  SKIP (ext {ext}): {f.name}")
            stats["skipped_ext"] += 1
            continue

        # Duplicate hash check
        try:
            fhash = file_md5(f)
        except Exception:
            fhash = None

        if fhash and fhash in seen_hashes:
            log.info(f"  SKIP (dup of {seen_hashes[fhash]}): {f.name}")
            stats["skipped_dup"] += 1
            continue

        if fhash:
            seen_hashes[fhash] = f.name

        # Extract
        vault_name = safe_stem(category, f.name)
        vault_path = VAULT / vault_name

        try:
            extractor = EXTRACTORS[ext]
            text = extractor(f)

            if not text or len(text.strip()) < 20:
                log_error(f.name, "Extracted text too short or empty")
                stats["errors"] += 1
                continue

            # Write to vault
            header = (
                f"=== SOURCE: {f.name} ===\n"
                f"=== CATEGORY: {category} ===\n"
                f"=== EXTRACTED: {datetime.now().isoformat()} ===\n"
                f"{'='*60}\n\n"
            )
            vault_path.write_text(header + text.strip(), encoding="utf-8")

            char_count = len(text.strip())
            word_count = len(text.split())
            log.info(f"  ✓ {f.name} → {vault_name} ({word_count:,} words)")
            stats["processed"] += 1
            stats["files"].append({
                "source": f.name,
                "vault": vault_name,
                "chars": char_count,
                "words": word_count,
                "category": category,
            })

        except Exception as e:
            log_error(f.name, str(e))
            stats["errors"] += 1

    return stats


def main():
    log.info("="*60)
    log.info("CAREER BRAIN — PHASE 1: NORMALIZE VAULT")
    log.info(f"Started: {datetime.now().isoformat()}")
    log.info("="*60)

    seen_hashes = {}  # Cross-directory duplicate detection
    all_stats = {}

    CATEGORIES = [
        ("resumes",       RAW_DOCS / "resumes"),
        ("cover_letters", RAW_DOCS / "cover_letters"),
        ("ksc",           RAW_DOCS / "ksc"),
        ("references",    RAW_DOCS / "references"),
        ("knowledge",     RAW_DOCS / "knowledge"),
    ]

    for category, src_dir in CATEGORIES:
        stats = process_directory(src_dir, category, seen_hashes)
        all_stats[category] = stats

    # Write error log
    if error_entries:
        ERROR_LOG.write_text("\n".join(error_entries) + "\n", encoding="utf-8")
        log.warning(f"\nParsing errors logged to: {ERROR_LOG}")

    # ─── HEALTH LEDGER ───────────────────────────────────────────────────────
    log.info("\n" + "="*60)
    log.info("PHASE 1 HEALTH LEDGER")
    log.info("="*60)

    total_processed = 0
    total_words = 0
    total_errors = 0

    for category, stats in all_stats.items():
        processed = stats["processed"]
        errors = stats["errors"]
        skipped_dup = stats["skipped_dup"]
        skipped_ext = stats["skipped_ext"]
        words = sum(f["words"] for f in stats["files"])
        total_processed += processed
        total_words += words
        total_errors += errors

        log.info(f"\n  [{category.upper()}]")
        log.info(f"    Extracted : {processed} files  ({words:,} words)")
        log.info(f"    Skipped   : {skipped_dup} duplicates, {skipped_ext} unsupported")
        log.info(f"    Errors    : {errors}")

    log.info("\n" + "─"*60)
    log.info(f"  TOTAL EXTRACTED : {total_processed} files")
    log.info(f"  TOTAL WORDS     : {total_words:,}")
    log.info(f"  TOTAL ERRORS    : {total_errors}")
    log.info(f"  VAULT LOCATION  : {VAULT}")
    log.info("="*60)

    # List all vault files
    vault_files = sorted(VAULT.glob("*.txt"))
    log.info(f"\nVault contains {len(vault_files)} .txt files:")
    for vf in vault_files:
        size = vf.stat().st_size
        log.info(f"  {vf.name}  ({size:,} bytes)")

    log.info("\nPhase 1 complete. Review health ledger, then approve Gate 3.")
    return total_errors == 0


if __name__ == "__main__":
    success = main()
    sys.exit(0 if success else 1)
